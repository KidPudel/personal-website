from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


# Standard business brief preset with one named override: ATS resume.
# The override keeps the same restrained Arial/black visual system while using
# resume-appropriate page density and no tables, columns, text boxes, or
# important information in headers/footers.
FONT = "Arial"
BLACK = RGBColor(0, 0, 0)
INK = RGBColor(38, 38, 38)
MUTED = RGBColor(68, 68, 68)


def set_font_family(r_fonts):
    for theme_attribute in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        r_fonts.attrib.pop(qn(theme_attribute), None)
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), FONT)
    r_fonts.set(qn("w:cs"), FONT)


def set_run_font(
    run,
    *,
    size: float,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor = INK,
    underline: bool | None = None,
):
    run.font.name = FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    set_font_family(r_fonts)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline
    return run


def set_spacing(paragraph, *, before: float = 0, after: float = 0, line: float = 1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def add_bottom_border(paragraph, *, color: str, size: str, space: str):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_hyperlink(paragraph, text: str, url: str, *, size: float = 9.25):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    set_font_family(fonts)
    run_properties.append(fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "444444")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    run_properties.append(underline)

    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(int(size * 2)))
    run_properties.append(size_element)
    complex_size = OxmlElement("w:szCs")
    complex_size.set(qn("w:val"), str(int(size * 2)))
    run_properties.append(complex_size)

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "360")
    indent.set(qn("w:hanging"), "240")
    paragraph_properties.append(indent)
    level.append(paragraph_properties)

    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    set_font_family(fonts)
    run_properties.append(fonts)
    level.append(run_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return num_id


def apply_num_id(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_properties = p_pr.find(qn("w:numPr"))
    if num_properties is None:
        num_properties = OxmlElement("w:numPr")
        p_pr.insert(0, num_properties)
    for old_child in list(num_properties):
        num_properties.remove(old_child)
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_properties.append(level)
    num_properties.append(number)


def configure_document(document: Document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.56)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""

    normal = document.styles["Normal"]
    normal.font.name = FONT
    set_font_family(normal._element.rPr.rFonts)
    normal.font.size = Pt(9.8)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.0

    style_specs = (
        ("Resume Summary", 10.0, False, False, 0.0, 1.0, 1.035, INK),
        ("Resume Skills", 9.5, False, False, 0.0, 0.7, 1.0, INK),
        ("Resume Entry", 9.85, True, False, 2.2, 0.35, 1.0, BLACK),
        ("Resume Context", 9.1, False, False, 0.0, 0.65, 1.0, INK),
        ("Resume Bullet", 9.5, False, False, 0.0, 0.8, 1.0, INK),
        ("Resume Education", 9.6, False, False, 0.0, 0.0, 1.0, INK),
        ("Resume Section", 10.3, True, False, 4.6, 2.0, 1.0, BLACK),
    )
    for name, size, bold, italic, before, after, line, color in style_specs:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = FONT
        set_font_family(style._element.rPr.rFonts)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = line

    document.styles["Resume Section"].paragraph_format.keep_with_next = True


def add_section_heading(document: Document, text: str):
    paragraph = document.add_paragraph(style="Resume Section")
    outline_level = OxmlElement("w:outlineLvl")
    outline_level.set(qn("w:val"), "0")
    paragraph._p.get_or_add_pPr().append(outline_level)
    set_run_font(paragraph.add_run(text), size=10.3, bold=True, color=BLACK)
    add_bottom_border(paragraph, color="A6A6A6", size="4", space="3")
    return paragraph


def add_entry(document: Document, title: str, date: str, context: str):
    heading = document.add_paragraph(style="Resume Entry")
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.tab_stops.add_tab_stop(Inches(7.14), WD_TAB_ALIGNMENT.RIGHT)
    set_run_font(heading.add_run(title), size=9.85, bold=True, color=BLACK)
    set_run_font(heading.add_run("\t"), size=9.85, bold=True, color=BLACK)
    set_run_font(heading.add_run(date), size=9.85, bold=True, color=BLACK)
    paragraph = document.add_paragraph(style="Resume Context")
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run(context), size=9.1)
    return paragraph


def add_bullet(document: Document, num_id: int, text: str):
    paragraph = document.add_paragraph(style="Resume Bullet")
    apply_num_id(paragraph, num_id)
    set_run_font(paragraph.add_run(text), size=9.5)
    return paragraph


def add_labeled_line(document: Document, label: str, text: str):
    paragraph = document.add_paragraph(style="Resume Skills")
    set_run_font(paragraph.add_run(f"{label}: "), size=9.5, bold=True)
    set_run_font(paragraph.add_run(text), size=9.5)
    return paragraph


def build_resume(output_path: Path):
    document = Document()
    configure_document(document)
    num_id = add_bullet_numbering(document)

    document.core_properties.title = "Igor Kupchinenko - Product Designer Resume"
    document.core_properties.subject = "Product design, UX/UI design, interaction design, and product development"
    document.core_properties.author = "Igor Kupchinenko"
    document.core_properties.keywords = (
        "product designer, product design, UX design, UI design, UX/UI design, interaction design, "
        "user research, information architecture, user flows, wireframing, prototyping, Figma, visual design"
    )
    document.core_properties.comments = "ATS-readable single-column resume"

    name = document.add_paragraph()
    set_spacing(name, after=0)
    set_run_font(name.add_run("IGOR KUPCHINENKO"), size=21.5, bold=True, color=BLACK)

    role = document.add_paragraph()
    set_spacing(role, after=1.6)
    set_run_font(role.add_run("PRODUCT DESIGNER"), size=12.4, bold=True, color=MUTED)

    contact_one = document.add_paragraph()
    set_spacing(contact_one, after=0.35)
    set_run_font(contact_one.add_run("Email: "), size=9.25, color=MUTED)
    add_hyperlink(contact_one, "i.kupchinenko@gmail.com", "mailto:i.kupchinenko@gmail.com")
    set_run_font(contact_one.add_run("  |  Portfolio: "), size=9.25, color=MUTED)
    add_hyperlink(
        contact_one,
        "kidpudel.github.io/personal-website",
        "https://kidpudel.github.io/personal-website/",
    )

    contact_two = document.add_paragraph()
    set_spacing(contact_two, after=1.1)
    set_run_font(contact_two.add_run("LinkedIn: "), size=9.25, color=MUTED)
    add_hyperlink(contact_two, "linkedin.com/in/iggydev", "https://www.linkedin.com/in/iggydev")
    set_run_font(contact_two.add_run("  |  GitHub: "), size=9.25, color=MUTED)
    add_hyperlink(contact_two, "github.com/KidPudel", "https://github.com/KidPudel")
    add_bottom_border(contact_two, color="222222", size="18", space="7")

    add_section_heading(document, "SUMMARY")
    summary = document.add_paragraph(style="Resume Summary")
    summary_text = (
        "Product designer with a software engineering background. I use research to understand what people are "
        "trying to do, turn that into product and interaction decisions, and carry the design into working software. "
        "Recent work includes a production food-ordering app, a solo macOS activity monitor, and a Telegram learning "
        "prototype."
    )
    set_run_font(summary.add_run(summary_text), size=10.0)

    add_section_heading(document, "SKILLS")
    add_labeled_line(
        document,
        "Product design",
        "User-centered design, user research, user interviews, contextual inquiry, research synthesis, problem "
        "framing, competitive analysis, information architecture, user flows, wireframing, UX design, UI design, "
        "interaction design, visual design, high-fidelity prototyping, usability testing, responsive and adaptive "
        "design, motion design",
    )
    add_labeled_line(
        document,
        "Tools and implementation",
        "Figma, Flutter, Vue, FastAPI, Go, Python, PostgreSQL, Unity; mobile, desktop, web, AI-assisted development, "
        "and code-based prototypes",
    )

    add_section_heading(document, "PRODUCT DESIGN PROJECTS")
    add_entry(
        document,
        "Product Designer | Observatory | Solo project",
        "Jul 2026 - Aug 2026",
        "Native macOS activity monitor released as version 0.1.1.",
    )
    add_bullet(
        document,
        num_id,
        "Conducted 5 exploratory interviews and synthesized 2 working personas. The central finding was that people "
        "approached performance through an application or task, not an isolated process.",
    )
    add_bullet(
        document,
        num_id,
        "Used that finding to make applications the main unit of the product, with repeatable test recordings, local "
        "saved results, and synchronized comparisons.",
    )
    add_bullet(
        document,
        num_id,
        "Designed the information architecture, Figma wireframes, interface, visual identity, and motion. Wrote "
        "detailed specifications, directed AI-assisted implementation, and shipped public macOS release 0.1.1 with a "
        "launch film.",
    )

    add_entry(
        document,
        "Product Designer | Two Sticks | Collaborative project",
        "May 2024 - Jun 2024",
        "Telegram Chinese character learning prototype based on a 2024 MPGU diploma team's research.",
    )
    add_bullet(
        document,
        num_id,
        "Turned the diploma team's pedagogical research into one Telegram learning loop: find, save, recall, study, "
        "and handwrite a Chinese character without leaving the chat.",
    )
    add_bullet(
        document,
        num_id,
        "Designed the information architecture, interaction, and interface, then built a working prototype across a "
        "Telegram bot, FastAPI backend, PostgreSQL data model, and Vue handwriting app with recognition scoring.",
    )

    add_section_heading(document, "WORK EXPERIENCE")
    add_entry(
        document,
        "Software Developer | Paycos | Contract",
        "Jul 2024 - Dec 2025",
        "Payments, order processing, and digital receipt systems.",
    )
    add_bullet(
        document,
        num_id,
        "Reduced PDF receipt processing by about 90%, from 6-10 seconds to 600-650 ms, by replacing unnecessary OCR "
        "with direct text extraction while preserving OCR for image inputs.",
    )
    add_bullet(
        document,
        num_id,
        "Built payment and order-processing workflows across Go, PostgreSQL, Kafka, Redis, gRPC, and Protobuf, "
        "translating product rules into stateful services.",
    )

    add_entry(
        document,
        "Product Designer / Mobile Engineer | PizzaSushiWok (SuperGood)",
        "Jan 2023 - Jul 2024",
        "Hired as a mobile engineer; role expanded into end-to-end product design ownership for a production "
        "food-ordering app.",
    )
    add_bullet(
        document,
        num_id,
        "Proposed and delivered a single Flutter app to replace separate Android and iOS applications, designing and "
        "implementing the end-to-end order experience from menu through checkout and delivery tracking.",
    )
    add_bullet(
        document,
        num_id,
        "Used customer conversations, contextual inquiry, competitive analysis, task tests, and polls to identify 3 "
        "recurring problems: unclear meal value, hidden delivery time and total, and friction in repeat orders.",
    )
    add_bullet(
        document,
        num_id,
        "Shipped fuller dish details, a persistent order summary, and direct paths from favorites and order history "
        "back to the basket across adaptive mobile layouts.",
    )

    add_entry(
        document,
        "Mobile Game Designer & Prototype Developer | 22bytes",
        "Nov 2022 - Jan 2023",
        "Short-term mobile game studio role.",
    )
    add_bullet(
        document,
        num_id,
        "Built playable Android prototypes in short cycles with Unity and Jetpack Compose, shaping the interaction "
        "flow, player feedback, presentation, and scope for each concept.",
    )

    add_section_heading(document, "EDUCATION")
    education = document.add_paragraph(style="Resume Education")
    education.paragraph_format.tab_stops.add_tab_stop(Inches(7.14), WD_TAB_ALIGNMENT.RIGHT)
    set_run_font(
        education.add_run("Bachelor of Computer & Information Science | Moscow Finance and Law Academy (MFUA)"),
        size=9.6,
    )
    set_run_font(education.add_run("\t"), size=9.6)
    set_run_font(education.add_run("2019 - 2023"), size=9.6)

    add_section_heading(document, "LANGUAGES")
    languages = document.add_paragraph(style="Resume Education")
    set_run_font(languages.add_run("Russian (native) | English (B2, upper-intermediate)"), size=9.6)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build_resume(args.output)


if __name__ == "__main__":
    main()
