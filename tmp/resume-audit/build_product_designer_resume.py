from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


FONT = "Arial"
INK = RGBColor(22, 22, 22)
MUTED = RGBColor(92, 92, 92)
LINK = RGBColor(34, 77, 124)
RULE = "C9C9C9"


def set_run_font(
    run,
    *,
    size: float,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor = INK,
    underline: bool | None = None,
):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline
    return run


def set_paragraph_spacing(
    paragraph,
    *,
    before: float = 0,
    after: float = 0,
    line: float = 1.0,
):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def add_bottom_border(paragraph, *, color: str = RULE, size: str = "5"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_hyperlink(paragraph, text: str, url: str, *, size: float = 9.25):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    run_properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), str(LINK))
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(int(size * 2)))
    run_properties.append(size_element)
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "270")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "270")
    indent.set(qn("w:hanging"), "210")
    paragraph_properties.append(indent)
    level.append(paragraph_properties)

    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    run_properties.append(fonts)
    level.append(run_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return num_id


def apply_num_id(paragraph, num_id: int):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    num_properties = paragraph_properties.find(qn("w:numPr"))
    if num_properties is None:
        num_properties = OxmlElement("w:numPr")
        paragraph_properties.insert(0, num_properties)
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_properties.append(level)
    num_properties.append(number)


def add_section_heading(document: Document, text: str):
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(text)
    add_bottom_border(paragraph)
    return paragraph


def add_entry_heading(document: Document, role_company: str, dates: str):
    paragraph = document.add_paragraph(style="Heading 2")
    fmt = paragraph.paragraph_format
    fmt.tab_stops.add_tab_stop(Inches(7.35), WD_TAB_ALIGNMENT.RIGHT)
    fmt.keep_with_next = True
    set_run_font(paragraph.add_run(role_company), size=10.0, bold=True)
    paragraph.add_run("\t")
    set_run_font(paragraph.add_run(dates), size=9.6, bold=True, color=MUTED)
    return paragraph


def add_resume_bullet(document: Document, num_id: int, text: str):
    paragraph = document.add_paragraph(style="Resume Bullet")
    apply_num_id(paragraph, num_id)
    set_run_font(paragraph.add_run(text), size=9.5)
    return paragraph


def add_labeled_line(document: Document, label: str, text: str):
    paragraph = document.add_paragraph(style="Resume Compact")
    set_run_font(paragraph.add_run(f"{label}: "), size=9.4, bold=True)
    set_run_font(paragraph.add_run(text), size=9.4)
    return paragraph


def configure_document(document: Document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.46)
    section.bottom_margin = Inches(0.43)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)
    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for name, size, bold, before, after, line in (
        ("Resume Bullet", 9.5, False, 0.0, 0.8, 1.0),
        ("Resume Compact", 9.4, False, 0.0, 0.5, 1.0),
        ("Resume Summary", 9.7, False, 0.0, 1.0, 1.04),
    ):
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = INK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line

    for name, size, before, after in (
        ("Heading 1", 10.5, 5.0, 2.2),
        ("Heading 2", 10.0, 2.5, 0.6),
    ):
        style = document.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = False
        style.font.color.rgb = INK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True


def build_resume(output_path: Path):
    document = Document()
    configure_document(document)
    num_id = add_bullet_numbering(document)

    document.core_properties.title = "Igor Kupchinenko - Product Designer Resume"
    document.core_properties.subject = "Product design, UX/UI design, interaction design, and product development"
    document.core_properties.author = "Igor Kupchinenko"
    document.core_properties.keywords = (
        "product designer, UX/UI design, interaction design, user research, "
        "prototyping, Figma, visual design, motion design"
    )
    document.core_properties.comments = "ATS-readable single-column resume"

    name = document.add_paragraph()
    set_paragraph_spacing(name, after=0)
    set_run_font(name.add_run("IGOR KUPCHINENKO"), size=22.0, bold=True)

    role = document.add_paragraph()
    set_paragraph_spacing(role, after=2.0)
    set_run_font(role.add_run("PRODUCT DESIGNER"), size=12.4, bold=True)

    contact_one = document.add_paragraph()
    set_paragraph_spacing(contact_one, after=0.4)
    add_hyperlink(contact_one, "i.kupchinenko@gmail.com", "mailto:i.kupchinenko@gmail.com")
    set_run_font(contact_one.add_run("  |  Portfolio: "), size=9.25, color=MUTED)
    add_hyperlink(
        contact_one,
        "kidpudel.github.io/personal-website",
        "https://kidpudel.github.io/personal-website/",
    )

    contact_two = document.add_paragraph()
    set_paragraph_spacing(contact_two, after=1.4)
    set_run_font(contact_two.add_run("LinkedIn: "), size=9.25, color=MUTED)
    add_hyperlink(contact_two, "linkedin.com/in/iggydev", "https://www.linkedin.com/in/iggydev")
    set_run_font(contact_two.add_run("  |  GitHub: "), size=9.25, color=MUTED)
    add_hyperlink(contact_two, "github.com/KidPudel", "https://github.com/KidPudel")

    add_section_heading(document, "SUMMARY")
    summary = document.add_paragraph(style="Resume Summary")
    set_run_font(
        summary.add_run(
            "Product designer with a software engineering and independent game development background. "
            "I find the real need behind a product, then carry the experience from user research and problem "
            "framing through interaction design, visual design, prototyping, and implementation. Game design "
            "informs how I use feedback, motion, pacing, and system behavior to make products clear, engaging, "
            "and a little more joyful."
        ),
        size=9.7,
    )

    add_section_heading(document, "SKILLS")
    add_labeled_line(
        document,
        "Product design",
        "User-centered design, product direction, user research, user interviews, problem framing, "
        "information architecture, user flows, wireframing, UX/UI design, interaction design, visual design, "
        "high-fidelity prototyping, motion design, usability testing, adaptive design",
    )
    add_labeled_line(
        document,
        "Tools and implementation",
        "Figma, Flutter, Godot, Unity, Vue, FastAPI, Go; mobile, desktop, web, and playable products",
    )

    add_section_heading(document, "PRODUCT DESIGN EXPERIENCE")
    add_entry_heading(document, "Independent Product Designer | Observatory", "Jul 2026 - Aug 2026")
    add_resume_bullet(
        document,
        num_id,
        "Conducted 5 exploratory user interviews and synthesized 2 working personas, revealing that people "
        "investigated the performance of an application or task, not an isolated process.",
    )
    add_resume_bullet(
        document,
        num_id,
        "Designed the product direction, information architecture, user flows, Figma wireframes, UX/UI, "
        "interaction design, visual identity, and motion for a native macOS activity monitor.",
    )
    add_resume_bullet(
        document,
        num_id,
        "Released version 0.1.1 with application-level process totals, controlled tests, and local, reopenable "
        "comparisons; directed AI-assisted implementation from authored specifications and produced the launch film.",
    )

    add_entry_heading(
        document,
        "Mobile Product Designer & Developer | PizzaSushiWok (SuperGood)",
        "Jan 2023 - Jul 2024",
    )
    add_resume_bullet(
        document,
        num_id,
        "Proposed replacing separate Android and iOS apps with one Flutter product, then owned product direction, "
        "user research, information architecture, UX/UI design, and implementation through production.",
    )
    add_resume_bullet(
        document,
        num_id,
        "Used customer conversations, contextual inquiry, competitive analysis, prototype task tests, comparison "
        "tests, and polls to reshape menu navigation, dish details, repeat ordering, checkout, payment and 3DS, "
        "maps, courier tracking, loyalty, and promotions.",
    )
    add_resume_bullet(
        document,
        num_id,
        "Designed adaptive layouts and implemented the end-to-end order journey, improving usability and "
        "reliability while requirements changed.",
    )

    add_entry_heading(document, "Product Designer & Developer | Two Sticks (Chinese Bee)", "2024")
    add_resume_bullet(
        document,
        num_id,
        "Translated a 2024 MPGU diploma team's pedagogical research into a coherent Telegram learning flow across "
        "search, saved vocabulary, flashcards, character guidance, notebook sheets, handwriting practice, and OCR scoring.",
    )
    add_resume_bullet(
        document,
        num_id,
        "Owned product direction, information architecture, UX/UI and interaction design, then built the bot, "
        "FastAPI backend, data model, and Vue web app into a shipped prototype.",
    )

    add_section_heading(document, "GAME DESIGN EXPERIENCE")
    add_entry_heading(
        document,
        "Mobile Game Designer & Prototype Developer | Short-term studio",
        "Nov 2022 - Jan 2023",
    )
    add_resume_bullet(
        document,
        num_id,
        "Turned mobile game concepts into playable Android prototypes in short cycles using Unity and Jetpack "
        "Compose, owning player experience, interaction flow, feedback, presentation, and scope.",
    )
    add_entry_heading(
        document,
        "Independent Game Designer & Developer | Self-directed projects",
        "2019 - Present",
    )
    add_resume_bullet(
        document,
        num_id,
        "Designed and shipped 2D and 3D Godot games, combining mechanics, progression, difficulty, feedback, "
        "visual art, animation, sound, and stateful systems into complete player experiences.",
    )
    add_resume_bullet(
        document,
        num_id,
        "Iterated Santa Foundation after external player feedback by changing hazard visibility, movement, spawning, "
        "timing, and difficulty; built Discourses by Campfire's environment, data-driven systems, 3D art, music, and sound.",
    )

    add_section_heading(document, "SOFTWARE ENGINEERING EXPERIENCE")
    add_entry_heading(document, "Software Developer | Paycos, contract", "Jul 2024 - Dec 2025")
    add_resume_bullet(
        document,
        num_id,
        "Developed payment, order-processing, and digital-receipt workflows in Go across PostgreSQL, Kafka, Redis, "
        "gRPC, and Protobuf, translating complex product rules into stateful systems.",
    )

    add_section_heading(document, "EDUCATION AND LANGUAGES")
    education = document.add_paragraph(style="Resume Compact")
    set_run_font(
        education.add_run(
            "Bachelor of Computer & Information Science | Moscow Finance and Law Academy (MFUA) | 2019 - 2023"
        ),
        size=9.4,
    )
    language = document.add_paragraph(style="Resume Compact")
    set_run_font(language.add_run("English: B2"), size=9.4)

    # Keep the final block together when it naturally fits, but allow Word to flow if it does not.
    education.paragraph_format.keep_with_next = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build_resume(args.output)


if __name__ == "__main__":
    main()
