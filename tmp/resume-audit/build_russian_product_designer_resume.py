from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches

from build_ats_product_designer_resume import (
    BLACK,
    MUTED,
    add_bottom_border,
    add_bullet,
    add_bullet_numbering,
    add_entry,
    add_hyperlink,
    add_labeled_line,
    add_section_heading,
    configure_document,
    set_run_font,
    set_spacing,
)


def build_resume(output_path: Path):
    document = Document()
    configure_document(document)
    num_id = add_bullet_numbering(document)

    document.core_properties.title = "Игорь Купчиненко - резюме продуктового дизайнера"
    document.core_properties.subject = (
        "Продуктовый дизайн, UX/UI-дизайн, дизайн взаимодействия и разработка продуктов"
    )
    document.core_properties.author = "Игорь Купчиненко"
    document.core_properties.keywords = (
        "продуктовый дизайнер, продуктовый дизайн, Product Designer, UX-дизайн, UI-дизайн, "
        "UX/UI-дизайн, дизайн взаимодействия, пользовательские исследования, информационная "
        "архитектура, пользовательские сценарии, прототипирование, Figma, визуальный дизайн"
    )
    document.core_properties.comments = "Одноколоночное ATS-совместимое резюме"

    name = document.add_paragraph()
    set_spacing(name, after=0)
    set_run_font(name.add_run("ИГОРЬ КУПЧИНЕНКО"), size=21.5, bold=True, color=BLACK)

    role = document.add_paragraph()
    set_spacing(role, after=1.6)
    set_run_font(role.add_run("ПРОДУКТОВЫЙ ДИЗАЙНЕР"), size=12.4, bold=True, color=MUTED)

    contact_one = document.add_paragraph()
    set_spacing(contact_one, after=0.35)
    set_run_font(contact_one.add_run("Email: "), size=9.25, color=MUTED)
    add_hyperlink(contact_one, "i.kupchinenko@gmail.com", "mailto:i.kupchinenko@gmail.com")
    set_run_font(contact_one.add_run("  |  Портфолио: "), size=9.25, color=MUTED)
    add_hyperlink(
        contact_one,
        "kidpudel.github.io/personal-website",
        "https://kidpudel.github.io/personal-website/",
    )

    contact_two = document.add_paragraph()
    set_spacing(contact_two, after=1.1)
    set_run_font(contact_two.add_run("LinkedIn: "), size=9.25, color=MUTED)
    add_hyperlink(contact_two, "linkedin.com/in/iggydev", "https://www.linkedin.com/in/iggydev")
    set_run_font(contact_two.add_run("  |  GitHub: "), size=9.25, color=MUTED)
    add_hyperlink(contact_two, "github.com/KidPudel", "https://github.com/KidPudel")
    add_bottom_border(contact_two, color="222222", size="18", space="7")

    add_section_heading(document, "ПРОФИЛЬ")
    summary = document.add_paragraph(style="Resume Summary")
    summary_text = (
        "Продуктовый дизайнер с опытом разработки ПО. Исследую задачи людей, перевожу выводы в "
        "продуктовые и интерфейсные решения и довожу дизайн до работающего продукта. Последние "
        "работы: приложение для заказа еды, собственный монитор активности для macOS и учебный "
        "прототип в Telegram."
    )
    set_run_font(summary.add_run(summary_text), size=10.0)

    add_section_heading(document, "НАВЫКИ")
    add_labeled_line(
        document,
        "Продуктовый дизайн",
        "Пользовательские исследования и интервью, контекстное исследование, синтез результатов, "
        "постановка проблемы, конкурентный анализ, информационная архитектура, пользовательские "
        "сценарии, вайрфреймы, UX/UI-дизайн, дизайн взаимодействия, визуальный дизайн, "
        "прототипирование, юзабилити-тестирование, адаптивный дизайн, моушн-дизайн",
    )
    add_labeled_line(
        document,
        "Инструменты и реализация",
        "Figma, Flutter, Vue, FastAPI, Go, Python, PostgreSQL, Unity; мобильные, десктопные и "
        "веб-продукты, AI-инструменты и прототипы в коде",
    )

    add_section_heading(document, "ПРОЕКТЫ В ПРОДУКТОВОМ ДИЗАЙНЕ")
    add_entry(
        document,
        "Продуктовый дизайнер | Observatory | Самостоятельный проект",
        "июль 2026 - август 2026",
        "Нативный монитор активности для macOS, выпущенный как версия 0.1.1.",
    )
    add_bullet(
        document,
        num_id,
        "Провел 5 исследовательских интервью и сформировал 2 рабочие персоны. Главный вывод: люди "
        "оценивали производительность через приложение или задачу, а не отдельный процесс.",
    )
    add_bullet(
        document,
        num_id,
        "На основе этого вывода сделал приложение основной единицей продукта и добавил повторяемые "
        "тестовые записи, локально сохраненные результаты и синхронизированные сравнения.",
    )
    add_bullet(
        document,
        num_id,
        "Спроектировал информационную архитектуру, вайрфреймы в Figma, интерфейс, визуальный стиль и "
        "анимацию. Подготовил спецификации, направлял реализацию с AI-инструментами и выпустил "
        "публичную версию 0.1.1 с релизным роликом.",
    )

    add_entry(
        document,
        "Продуктовый дизайнер | Two Sticks | Командный проект",
        "май 2024 - июнь 2024",
        "Прототип для изучения китайских иероглифов в Telegram на основе исследования дипломной "
        "команды МПГУ 2024 года.",
    )
    add_bullet(
        document,
        num_id,
        "Перевел педагогическое исследование в единый сценарий внутри Telegram: найти, сохранить, "
        "вспомнить, изучить и написать китайский иероглиф, не выходя из чата.",
    )
    add_bullet(
        document,
        num_id,
        "Спроектировал информационную архитектуру, взаимодействие и интерфейс, затем собрал рабочий "
        "прототип: Telegram-бот, бэкенд на FastAPI, модель данных PostgreSQL и веб-приложение для "
        "письма на Vue с оценкой распознавания.",
    )

    add_section_heading(document, "ОПЫТ РАБОТЫ")
    add_entry(
        document,
        "Разработчик ПО | Paycos | Контракт",
        "июль 2024 - декабрь 2025",
        "Платежи, обработка заказов и системы электронных чеков.",
    )
    add_bullet(
        document,
        num_id,
        "Ускорил обработку PDF-чеков примерно на 90%: с 6-10 секунд до 600-650 мс. Заменил "
        "ненужный OCR прямым извлечением текста, сохранив OCR для изображений.",
    )
    add_bullet(
        document,
        num_id,
        "Разрабатывал процессы оплаты и обработки заказов на Go, PostgreSQL, Kafka, Redis, gRPC и "
        "Protobuf, переводя продуктовые правила в сервисы с сохранением состояния.",
    )

    add_entry(
        document,
        "Продуктовый дизайнер / мобильный разработчик | PizzaSushiWok (SuperGood)",
        "январь 2023 - июль 2024",
        "Пришел как мобильный разработчик и в процессе взял на себя продуктовый дизайн всего "
        "приложения для заказа еды.",
    )
    add_bullet(
        document,
        num_id,
        "Предложил и выпустил единое Flutter-приложение вместо отдельных версий для Android и iOS. "
        "Спроектировал и реализовал весь сценарий заказа - от меню до оплаты и отслеживания доставки.",
    )
    add_bullet(
        document,
        num_id,
        "Использовал беседы с клиентами, контекстные исследования, конкурентный анализ, тесты "
        "сценариев и опросы, чтобы выявить 3 проблемы: непонятную ценность блюд, скрытые сроки "
        "доставки и итоговую сумму, а также неудобный повторный заказ.",
    )
    add_bullet(
        document,
        num_id,
        "Выпустил подробные карточки блюд, закрепленную сводку заказа и быстрый переход из избранного "
        "и истории заказов обратно в корзину; адаптировал интерфейс под разные мобильные экраны.",
    )

    add_entry(
        document,
        "Гейм-дизайнер мобильных игр и разработчик прототипов | 22bytes",
        "ноябрь 2022 - январь 2023",
        "Краткосрочная работа в студии мобильных игр.",
    )
    add_bullet(
        document,
        num_id,
        "В коротких циклах собирал игровые Android-прототипы на Unity и Jetpack Compose; отвечал за "
        "сценарий взаимодействия, обратную связь, подачу и объем каждой концепции.",
    )

    add_section_heading(document, "ОБРАЗОВАНИЕ И ЯЗЫКИ")
    education = document.add_paragraph(style="Resume Education")
    education.paragraph_format.tab_stops.add_tab_stop(Inches(7.14), WD_TAB_ALIGNMENT.RIGHT)
    set_run_font(
        education.add_run("Бакалавр компьютерных и информационных наук | МФЮА"),
        size=9.6,
    )
    set_run_font(education.add_run("\t"), size=9.6)
    set_run_font(education.add_run("2019 - 2023"), size=9.6)

    languages = document.add_paragraph(style="Resume Education")
    set_spacing(languages, before=0.8)
    set_run_font(languages.add_run("Русский (родной) | Английский (B2, Upper-Intermediate)"), size=9.6)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build_resume(args.output)


if __name__ == "__main__":
    main()
