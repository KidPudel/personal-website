from docx import Document


SOURCE = "public/resume/igor-kupchinenko-resume.docx"
OUTPUT = "tmp/pdfs/igor-kupchinenko-resume.docx"

REPLACEMENTS = {
    "EXPERIENCE DESIGN & INTERACTIVE PRODUCT DEVELOPMENT | PRODUCT DESIGN ENGINEER":
        "PRODUCT & INTERACTION DESIGN | INTERACTIVE PRODUCT DEVELOPMENT",
    (
        "I design interactive product experiences and build or direct the systems that make them real. "
        "I work from the user's desired capability through product shape, interaction flow, feedback, "
        "state, rules, interface, backend behavior, data, integrations, and delivery. The strongest "
        "evidence is repeated ownership of the whole experience: a released native macOS analysis "
        "product created from idea through launch film, a production food-ordering product, an "
        "independent learning product, stateful fintech workflows, and completed games shaped through "
        "player feedback."
    ): (
        "I design interactive product experiences and use prototypes, motion, code, and AI-directed "
        "development to make them real. I start with what a person should be able to do and how the "
        "product should feel and behave, then choose the tools needed to carry that experience into a "
        "working artifact. My strongest evidence is repeated ownership of whole experiences: a native "
        "macOS analysis product created from idea through launch film, a production food-ordering "
        "product, an independent learning product, complex payment workflows, and completed games "
        "shaped through player feedback."
    ),
}


def replace_paragraph_text(paragraph, old, new):
    if paragraph.text != old:
        return False

    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new)
        return True

    runs[0].text = new
    for run in runs[1:]:
        run.text = ""
    return True


document = Document(SOURCE)
found = {old: False for old in REPLACEMENTS}

for paragraph in document.paragraphs:
    for old, new in REPLACEMENTS.items():
        if replace_paragraph_text(paragraph, old, new):
            found[old] = True

missing = [old for old, was_found in found.items() if not was_found]
if missing:
    raise RuntimeError(f"Expected resume text not found: {missing}")

document.save(OUTPUT)
